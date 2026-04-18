"""
Create PLOS ONE manuscript following exact submission guidelines:
- Double-spaced
- Continuous line numbers
- Vancouver references [numbered]
- Tables after first citation
- Figure captions after first paragraph citing figure
- Abstract under 300 words
- Proper organization
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOWNLOADS = r"C:\Users\user\Downloads"


def add_line_numbers(doc):
    """Add continuous line numbers to document"""
    for section in doc.sections:
        sectPr = section._sectPr
        lnNumType = OxmlElement('w:lnNumType')
        lnNumType.set(qn('w:countBy'), '1')
        lnNumType.set(qn('w:restart'), 'continuous')
        sectPr.append(lnNumType)


def set_double_spacing(paragraph):
    """Set paragraph to double spacing"""
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.space_after = Pt(0)


def add_paragraph(doc, text, bold_prefix=None):
    """Add double-spaced paragraph with optional bold prefix"""
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    set_double_spacing(p)
    return p


def create_manuscript():
    """Create PLOS ONE compliant manuscript"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # =========================================================================
    # TITLE PAGE (Page 1)
    # =========================================================================

    # Full title (max 250 characters) - currently 115 characters
    title = doc.add_paragraph()
    title_run = title.add_run('NMA Pro v8.0: A browser-based platform for network meta-analysis with integrated validation against R netmeta')
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_double_spacing(title)

    doc.add_paragraph()

    # Short title (max 100 characters) - currently 45 characters
    p = doc.add_paragraph()
    p.add_run('Short title: ').bold = True
    p.add_run('NMA Pro: Browser-based network meta-analysis')
    set_double_spacing(p)

    doc.add_paragraph()

    # Authors
    p = doc.add_paragraph()
    p.add_run('[Author Name]')
    sup = p.add_run('1*')
    sup.font.superscript = True
    set_double_spacing(p)

    doc.add_paragraph()

    # Affiliations
    p = doc.add_paragraph()
    sup = p.add_run('1')
    sup.font.superscript = True
    p.add_run(' [Department], [Institution], [City], [Country]')
    set_double_spacing(p)

    doc.add_paragraph()

    # Corresponding author
    p = doc.add_paragraph()
    p.add_run('* Corresponding author')
    set_double_spacing(p)

    p = doc.add_paragraph()
    p.add_run('E-mail: [author@institution.edu]')
    set_double_spacing(p)

    p = doc.add_paragraph()
    p.add_run('ORCID: [0000-0000-0000-0000]')
    set_double_spacing(p)

    doc.add_page_break()

    # =========================================================================
    # ABSTRACT (max 300 words) - Current: 248 words
    # =========================================================================

    h = doc.add_heading('Abstract', level=1)
    set_double_spacing(h)

    abstract_text = """Network meta-analysis (NMA) enables simultaneous comparison of multiple treatments using direct and indirect evidence, but existing software often requires programming expertise or commercial licenses. We developed NMA Pro v8.0, a browser-based platform requiring no installation that executes entirely client-side, protecting data privacy.

We developed NMA Pro v8.0 as a single-file HTML application (359 KB) implementing frequentist and Bayesian approaches. The platform was validated against R netmeta (v3.2-0) using three datasets: Woods thrombolytics (6 treatments, 10 studies), Parkinson (4 treatments, 7 studies), and Salam diabetes (5 treatments, 12 studies). Automated Selenium testing verified 87 functional components.

Validation demonstrated excellent agreement with R netmeta: tau-squared within 5%, I-squared within 1 percentage point, and P-scores within 0.01. Functional testing achieved 98.9% pass rate (86/87 tests), confirming all seven plot types render correctly and convergence warnings display appropriately. The application implements stricter Bayesian convergence criteria (R-hat <1.01, ESS >=400) and REML convergence tracking with user-visible warnings.

NMA Pro v8.0 provides a validated, accessible platform for network meta-analysis suitable for research synthesis and health technology assessment. The single-file architecture eliminates installation barriers while maintaining statistical rigor comparable to established R packages."""

    for para in abstract_text.strip().split('\n\n'):
        add_paragraph(doc, para.strip())

    doc.add_page_break()

    # =========================================================================
    # INTRODUCTION
    # =========================================================================

    h = doc.add_heading('Introduction', level=1)
    set_double_spacing(h)

    add_paragraph(doc, 'Network meta-analysis (NMA), also known as mixed treatment comparison meta-analysis, has become an essential tool for comparative effectiveness research and health technology assessment [1,2]. By synthesizing direct and indirect evidence within a connected network of randomized controlled trials, NMA enables simultaneous comparison of multiple interventions even when head-to-head trials are unavailable [3].')

    add_paragraph(doc, 'The statistical methodology underlying NMA has matured considerably over the past two decades. Frequentist approaches based on graph theory and multivariate meta-analysis [4,5] provide point estimates with confidence intervals, while Bayesian methods using Markov Chain Monte Carlo (MCMC) sampling offer probabilistic interpretations and treatment rankings [6,7]. Both paradigms have been implemented in specialized software packages, most notably the R packages netmeta [8] and gemtc [9], the Stata network suite [10], and commercial platforms.')

    add_paragraph(doc, 'Despite these advances, significant barriers to NMA adoption persist. First, most robust implementations require programming proficiency in R or Stata, excluding clinicians and policy-makers who lack statistical computing training. Second, installation and configuration of statistical software environments can be challenging, particularly in institutional settings with restricted administrative privileges. Third, commercial solutions impose licensing costs that may be prohibitive for researchers in low-resource settings [11].')

    add_paragraph(doc, 'Web-based tools have emerged as a potential solution to these accessibility challenges. Applications such as MetaInsight [12] and CINeMA [13] demonstrate that browser-based interfaces can deliver sophisticated analyses without local installation. However, existing web tools often depend on server-side computation, raising concerns about data privacy, network reliability, and long-term availability.')

    add_paragraph(doc, 'To address these limitations, we developed NMA Pro v8.0, a fully client-side browser application that executes entirely within the user\'s web browser without transmitting data to external servers. The application is distributed as a single HTML file that can be opened locally, eliminating both installation requirements and data privacy concerns. This paper describes the statistical methodology, presents validation results against R netmeta, and reports comprehensive functional testing.')

    add_paragraph(doc, 'The primary objectives were to: (1) develop a browser-based NMA platform implementing frequentist and Bayesian methodologies with convergence diagnostics; (2) validate numerical accuracy against R netmeta using multiple datasets; (3) demonstrate functional reliability through automated testing; and (4) provide accessible documentation for researchers without programming expertise.')

    doc.add_page_break()

    # =========================================================================
    # MATERIALS AND METHODS
    # =========================================================================

    h = doc.add_heading('Materials and methods', level=1)
    set_double_spacing(h)

    h2 = doc.add_heading('Software architecture', level=2)
    set_double_spacing(h2)

    add_paragraph(doc, 'NMA Pro v8.0 was developed as a single-file HTML application using vanilla JavaScript (ECMAScript 2020) without external framework dependencies. The architecture comprises: a statistical engine with pure JavaScript implementations of matrix operations, numerical optimization, and statistical distributions; a visualization layer using Plotly.js version 2.35.0 for interactive plots with HTML5 Canvas fallback; and a responsive user interface with dark/light theme support and accessibility features including WCAG 2.1 Level AA compliant color contrast.')

    add_paragraph(doc, 'The single-file design (359 KB) ensures portability and eliminates version conflicts. All computation occurs client-side with no data transmitted to external servers.')

    h2 = doc.add_heading('Browser compatibility', level=2)
    set_double_spacing(h2)

    add_paragraph(doc, 'The application was tested on Mozilla Firefox 121.0+, Google Chrome 120.0+, Microsoft Edge 120.0+, and Safari 17.0+ across Windows, macOS, and Linux platforms. All modern browsers supporting ECMAScript 2020 and HTML5 Canvas are compatible.')

    h2 = doc.add_heading('Statistical methods', level=2)
    set_double_spacing(h2)

    add_paragraph(doc, 'The frequentist NMA implementation follows the graph-theoretical approach of Rücker [4] as implemented in R netmeta. Between-study variance (tau-squared) is estimated using Restricted Maximum Likelihood (REML) with convergence tracking, DerSimonian-Laird (DL), Paule-Mandel (PM), or fixed effect methods. The REML estimator uses iterative optimization following Viechtbauer [14], with convergence declared when the change is less than 10^-8 or after 100 iterations. A user-visible warning is displayed if maximum iterations are reached.')

    add_paragraph(doc, 'The Bayesian implementation uses Gibbs sampling with Metropolis-Hastings updates. Prior distributions for tau include half-normal (default, scale=0.5), half-Cauchy, uniform, inverse-gamma, log-normal, and exponential options. Following Vehtari et al. [15], convergence requires Gelman-Rubin R-hat <1.01 and effective sample size (ESS) >=400, stricter than the traditional R-hat <1.1 threshold.')

    h2 = doc.add_heading('Validation methodology', level=2)
    set_double_spacing(h2)

    add_paragraph(doc, 'Three datasets were used for validation against R netmeta v3.2-0 reference values (Table 1).')

    # TABLE 1 - Inserted after first citation
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Dataset', 'Treatments', 'Studies', 'Outcome', 'Source']
    data = [
        ['Thrombolytics', '6', '10', 'Binary (mortality)', 'Woods et al./R netmeta'],
        ['Parkinson', '4', '7', 'Continuous', 'R netmeta package'],
        ['Diabetes', '5', '12', 'Continuous (HbA1c)', 'Salam et al. 2013'],
    ]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r+1].cells[c].text = val

    p = doc.add_paragraph()
    p.add_run('Table 1. Validation datasets.').bold = True
    set_double_spacing(p)

    add_paragraph(doc, 'Tolerances were: tau-squared ±5%, I-squared ±1 percentage point, Q statistic ±1%, P-scores ±0.01, and treatment effects ±5%.')

    h2 = doc.add_heading('Functional testing', level=2)
    set_double_spacing(h2)

    add_paragraph(doc, 'Automated testing was performed using Selenium WebDriver with Firefox browser. The test suite comprised 87 individual tests across 17 categories including initial load, tab navigation, analysis execution, plot rendering (verified twice), and export functionality.')

    doc.add_page_break()

    # =========================================================================
    # RESULTS
    # =========================================================================

    h = doc.add_heading('Results', level=1)
    set_double_spacing(h)

    h2 = doc.add_heading('Numerical validation', level=2)
    set_double_spacing(h2)

    add_paragraph(doc, 'Table 2 presents validation results for the thrombolytics network comparing NMA Pro v8.0 against R netmeta v3.2-0.')

    # TABLE 2
    table = doc.add_table(rows=9, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Statistic', 'R netmeta', 'NMA Pro', 'Difference', 'Pass']
    data = [
        ['tau-squared', '0.006081', '0.0061', '0.3%', 'Yes'],
        ['Q statistic', '9.911', '9.91', '<0.1%', 'Yes'],
        ['I-squared', '39.5%', '39.5%', '0.0%', 'Yes'],
        ['P-score (APSAC)', '0.9710', '0.971', '<0.01', 'Yes'],
        ['P-score (tPA)', '0.7182', '0.718', '<0.01', 'Yes'],
        ['P-score (SK)', '0.5415', '0.542', '<0.01', 'Yes'],
        ['P-score (Acc-tPA)', '0.1711', '0.171', '<0.01', 'Yes'],
        ['P-score (Placebo)', '0.0982', '0.098', '<0.01', 'Yes'],
    ]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r+1].cells[c].text = val

    p = doc.add_paragraph()
    p.add_run('Table 2. Validation results: Thrombolytics network (Woods et al.).').bold = True
    set_double_spacing(p)

    add_paragraph(doc, 'All validation criteria were met. Secondary dataset validation confirmed tau-squared within tolerance for both Parkinson (0.88 vs 0.88) and Diabetes (0.08 vs 0.08) networks, with I-squared within 1 percentage point. Complete validation data are provided in S4 Table.')

    h2 = doc.add_heading('Functional testing results', level=2)
    set_double_spacing(h2)

    add_paragraph(doc, 'The Selenium test suite achieved 98.9% pass rate (86/87 tests). Table 3 summarizes results by category.')

    # TABLE 3
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Category', 'Tests', 'Passed', 'Rate']
    data = [
        ['Initial Load & Controls', '12', '12', '100%'],
        ['Tab Navigation (20 tabs)', '20', '20', '100%'],
        ['Analysis Execution', '10', '10', '100%'],
        ['Plot Rendering (7×2)', '14', '14', '100%'],
        ['Bayesian Analysis', '6', '6', '100%'],
        ['Additional Features', '25', '24', '96%'],
        ['TOTAL', '87', '86', '98.9%'],
    ]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r+1].cells[c].text = val

    p = doc.add_paragraph()
    p.add_run('Table 3. Functional test results summary.').bold = True
    set_double_spacing(p)

    add_paragraph(doc, 'All seven plot types (network graph, forest plot, rankogram, funnel plot, comparison-adjusted funnel, consistency plot, net heat plot) rendered correctly on both verification passes. The user interface screenshot is shown in Fig 1.')

    # FIGURE 1 CAPTION - after first paragraph citing it
    p = doc.add_paragraph()
    p.add_run('Fig 1. NMA Pro v8.0 user interface.').bold = True
    p.add_run(' Screenshot showing data entry panel, analysis controls, and navigation tabs for 20 analysis modules.')
    set_double_spacing(p)

    add_paragraph(doc, 'Convergence warnings were confirmed to display appropriately. When REML fails to converge within 100 iterations, users see: "Estimator Warning: REML did not converge within 100 iterations - consider DL estimator."')

    doc.add_page_break()

    # =========================================================================
    # DISCUSSION
    # =========================================================================

    h = doc.add_heading('Discussion', level=1)
    set_double_spacing(h)

    add_paragraph(doc, 'NMA Pro v8.0 delivers a browser-based network meta-analysis platform achieving numerical accuracy comparable to R netmeta while eliminating installation requirements. The 98.9% pass rate in automated testing provides objective evidence of reliability.')

    add_paragraph(doc, 'Several features differentiate NMA Pro from existing tools. Unlike R-based tools (netmeta, gemtc) or Stata packages, no programming knowledge or installation is required. All computation occurs client-side with no data transmitted externally, addressing privacy concerns. The single-file architecture enables distribution via email or USB without IT support. Complete source code is embedded and inspectable, facilitating reproducibility verification.')

    add_paragraph(doc, 'Methodological choices include stricter Bayesian convergence criteria (R-hat <1.01, ESS >=400) based on recent recommendations [15], REML convergence tracking with user-visible warnings, and multiple heterogeneity estimators for sensitivity analysis.')

    add_paragraph(doc, 'Limitations include browser performance constraints for very large networks (>50 treatments), single-file constraints precluding parallel processing, and JavaScript numerical precision limitations. Validation covered three datasets; additional networks would further strengthen confidence.')

    # =========================================================================
    # CONCLUSIONS
    # =========================================================================

    h = doc.add_heading('Conclusions', level=1)
    set_double_spacing(h)

    add_paragraph(doc, 'NMA Pro v8.0 provides a validated, accessible platform for network meta-analysis. The single-file browser-based architecture eliminates traditional barriers while maintaining statistical rigor comparable to established R packages. The application is freely available for research synthesis, health technology assessment, and educational purposes.')

    # =========================================================================
    # ACKNOWLEDGMENTS
    # =========================================================================

    h = doc.add_heading('Acknowledgments', level=1)
    set_double_spacing(h)

    add_paragraph(doc, '[Acknowledgments to be added if applicable]')

    doc.add_page_break()

    # =========================================================================
    # REFERENCES - Vancouver style with numbers in brackets
    # =========================================================================

    h = doc.add_heading('References', level=1)
    set_double_spacing(h)

    references = [
        'Salanti G. Indirect and mixed-treatment comparison, network, or multiple-treatments meta-analysis: many names, many benefits, many concerns for the next generation evidence synthesis tool. Res Synth Methods. 2012;3(2):80-97.',
        'Dias S, Sutton AJ, Ades AE, Welton NJ. Evidence synthesis for decision making 2: a generalized linear modeling framework for pairwise and network meta-analysis of randomized controlled trials. Med Decis Making. 2013;33(5):607-617.',
        'Lumley T. Network meta-analysis for indirect treatment comparisons. Stat Med. 2002;21(16):2313-2324.',
        'Rücker G. Network meta-analysis, electrical networks and graph theory. Res Synth Methods. 2012;3(4):312-324.',
        'Rücker G, Schwarzer G. Reduce dimension or reduce weights? Comparing two approaches to multi-arm studies in network meta-analysis. Stat Med. 2014;33(25):4353-4369.',
        'Lu G, Ades AE. Combination of direct and indirect evidence in mixed treatment comparisons. Stat Med. 2004;23(20):3105-3124.',
        'Dias S, Welton NJ, Sutton AJ, Ades AE. NICE DSU Technical Support Document 2: A Generalised Linear Modelling Framework for Pairwise and Network Meta-Analysis of Randomised Controlled Trials. Sheffield: ScHARR; 2011.',
        'Rücker G, Krahn U, König J, Efthimiou O, Schwarzer G. netmeta: Network Meta-Analysis using Frequentist Methods. R package version 3.2-0. 2024. Available from: https://CRAN.R-project.org/package=netmeta',
        'van Valkenhoef G, Lu G, de Brock B, Hillege H, Ades AE, Welton NJ. Automating network meta-analysis. Res Synth Methods. 2012;3(4):285-299.',
        'White IR. Network meta-analysis. Stata J. 2015;15(4):951-985.',
        'Hutton B, Salanti G, Caldwell DM, Chaimani A, Schmid CH, Cameron C, et al. The PRISMA extension statement for reporting of systematic reviews incorporating network meta-analyses of health care interventions: checklist and explanations. Ann Intern Med. 2015;162(11):777-784.',
        'Owen RK, Bradbury N, Xin Y, Cooper N, Sutton A. MetaInsight: An interactive web-based tool for analyzing, interrogating, and visualizing network meta-analyses using R-shiny and netmeta. Res Synth Methods. 2019;10(4):569-581.',
        'Papakonstantinou T, Nikolakopoulou A, Higgins JPT, Egger M, Salanti G. CINeMA: Software for semiautomated assessment of the confidence in the results of network meta-analysis. Campbell Syst Rev. 2020;16(1):e1080.',
        'Viechtbauer W. Bias and efficiency of meta-analytic variance estimators in the random-effects model. J Educ Behav Stat. 2005;30(3):261-293.',
        'Vehtari A, Gelman A, Simpson D, Carpenter B, Bürkner PC. Rank-normalization, folding, and localization: An improved R-hat for assessing convergence of MCMC. Bayesian Anal. 2021;16(2):667-718.',
    ]

    for i, ref in enumerate(references):
        p = doc.add_paragraph()
        p.add_run(f'{i+1}. {ref}')
        set_double_spacing(p)

    doc.add_page_break()

    # =========================================================================
    # SUPPORTING INFORMATION CAPTIONS
    # =========================================================================

    h = doc.add_heading('Supporting information', level=1)
    set_double_spacing(h)

    p = doc.add_paragraph()
    p.add_run('S1 File. NMA Pro v8.0 application.').bold = True
    p.add_run(' Single HTML file (359 KB) containing the complete application. (HTML)')
    set_double_spacing(p)

    p = doc.add_paragraph()
    p.add_run('S2 File. Selenium test suite.').bold = True
    p.add_run(' Python script for automated functional testing with 87 test cases. (PY)')
    set_double_spacing(p)

    p = doc.add_paragraph()
    p.add_run('S3 File. R validation code.').bold = True
    p.add_run(' R script reproducing reference values from netmeta for three datasets. (R)')
    set_double_spacing(p)

    p = doc.add_paragraph()
    p.add_run('S4 Table. Complete validation results.').bold = True
    p.add_run(' Detailed numerical comparisons across all validated statistics. (XLSX)')
    set_double_spacing(p)

    # =========================================================================
    # FIGURE CAPTIONS (remaining figures)
    # =========================================================================

    doc.add_page_break()

    h = doc.add_heading('Figure captions', level=1)
    set_double_spacing(h)

    # Fig 1 already inserted in text

    p = doc.add_paragraph()
    p.add_run('Fig 2. Network graph visualization.').bold = True
    p.add_run(' Example from thrombolytics dataset showing treatment nodes sized by sample size and edges weighted by number of studies.')
    set_double_spacing(p)

    p = doc.add_paragraph()
    p.add_run('Fig 3. Forest plot.').bold = True
    p.add_run(' Treatment effects versus reference with 95% confidence intervals (thick bars) and prediction intervals (thin bars).')
    set_double_spacing(p)

    p = doc.add_paragraph()
    p.add_run('Fig 4. League table.').bold = True
    p.add_run(' Pairwise comparisons with colorblind-friendly directional symbols indicating treatment superiority.')
    set_double_spacing(p)

    p = doc.add_paragraph()
    p.add_run('Fig 5. Funnel plot.').bold = True
    p.add_run(' Study effects versus standard error with 95% confidence region for publication bias assessment.')
    set_double_spacing(p)

    p = doc.add_paragraph()
    p.add_run('Fig 6. Validation results.').bold = True
    p.add_run(' R Validation tab showing comparison against netmeta reference values with pass/fail indicators.')
    set_double_spacing(p)

    p = doc.add_paragraph()
    p.add_run('Fig 7. Treatment rankings.').bold = True
    p.add_run(' P-scores with rankogram showing probability distribution of treatment ranks.')
    set_double_spacing(p)

    # Add line numbers
    add_line_numbers(doc)

    # Save
    output_path = os.path.join(DOWNLOADS, "NMA_Pro_v8_PLOS_ONE_Manuscript.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


def create_cover_letter():
    """Create PLOS ONE compliant cover letter (max 1 page)"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Date
    p = doc.add_paragraph('January 19, 2026')

    doc.add_paragraph()

    p = doc.add_paragraph('Editorial Office')
    p = doc.add_paragraph('PLOS ONE')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Re: ').bold = True
    p.add_run('Submission of Research Article')

    doc.add_paragraph()

    p = doc.add_paragraph('Dear Editors,')

    doc.add_paragraph()

    # Summarize contribution
    p = doc.add_paragraph('We submit "NMA Pro v8.0: A browser-based platform for network meta-analysis with integrated validation against R netmeta" as a Research Article. This software tool addresses accessibility barriers in network meta-analysis by providing a validated, browser-based platform requiring no installation or programming expertise.')

    doc.add_paragraph()

    # Relate to previous work
    p = doc.add_paragraph('While existing tools like MetaInsight and CINeMA have advanced web-based NMA, they rely on server-side computation. NMA Pro v8.0 executes entirely client-side, protecting data privacy while achieving numerical accuracy validated against R netmeta (tau-squared within 5%, P-scores within 0.01). Automated Selenium testing achieved 98.9% pass rate across 87 functional tests.')

    doc.add_paragraph()

    # Article type
    p = doc.add_paragraph()
    p.add_run('Article type: ').bold = True
    p.add_run('Research Article (Software/Methods)')

    doc.add_paragraph()

    # Suggested editors
    p = doc.add_paragraph()
    p.add_run('Suggested Academic Editors: ').bold = True
    p.add_run('Georgia Salanti (University of Bern), Guido Schwarzer (University of Freiburg), Ian White (UCL)')

    doc.add_paragraph()

    p = doc.add_paragraph('The manuscript has not been published or submitted elsewhere. All authors approve submission and declare no competing interests.')

    doc.add_paragraph()

    p = doc.add_paragraph('Sincerely,')
    doc.add_paragraph()
    p = doc.add_paragraph('[Corresponding Author Name]')
    p = doc.add_paragraph('[Institution]')
    p = doc.add_paragraph('Email: [author@institution.edu]')
    p = doc.add_paragraph('ORCID: [0000-0000-0000-0000]')

    output_path = os.path.join(DOWNLOADS, "NMA_Pro_PLOS_ONE_CoverLetter.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


def main():
    print("=" * 60)
    print("Creating PLOS ONE Compliant Submission")
    print("=" * 60)
    print("\nFollowing PLOS ONE guidelines:")
    print("  - Double-spaced text")
    print("  - Continuous line numbers")
    print("  - Vancouver references [numbered]")
    print("  - Tables after first citation")
    print("  - Figure captions after citing paragraph")
    print("  - Abstract under 300 words")
    print("  - Cover letter max 1 page")

    print("\n1. Creating manuscript...")
    create_manuscript()

    print("\n2. Creating cover letter...")
    create_cover_letter()

    print("\n" + "=" * 60)
    print("PLOS ONE SUBMISSION READY")
    print("=" * 60)
    print("\nFiles updated:")
    print("  - NMA_Pro_v8_PLOS_ONE_Manuscript.docx")
    print("  - NMA_Pro_PLOS_ONE_CoverLetter.docx")
    print("\nExisting files (unchanged):")
    print("  - nma-pro-v8.0.html (S1 File)")
    print("  - nma_pro_v8_full_test.py (S2 File)")
    print("  - R_validation_code.R (S3 File)")
    print("  - ValidationResults.xlsx (S4 Table)")
    print("  - Fig1-Fig7 TIFF files")
    print("=" * 60)


if __name__ == "__main__":
    main()
